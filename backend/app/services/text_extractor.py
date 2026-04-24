from typing import List, Tuple

import fitz
from docx import Document as DocxDocument

from ..core.logging_config import logger


class TextExtractor:
    @staticmethod
    def extract_pdf(pdf_path: str) -> List[Tuple[str, int | None]]:
        try:
            document = fitz.open(pdf_path)
            pages: List[Tuple[str, int | None]] = []
            for page_index in range(len(document)):
                page = document.load_page(page_index)
                text = page.get_text("text") or ""
                pages.append((text.strip(), page_index + 1))
            document.close()
            logger.info("Extracted %s pages from PDF %s", len(pages), pdf_path)
            return pages
        except Exception:
            logger.exception("PDF extraction failed for %s", pdf_path)
            raise

    @staticmethod
    def extract_docx(docx_path: str) -> List[Tuple[str, int | None]]:
        try:
            doc = DocxDocument(docx_path)
            content_blocks: List[str] = []

            for paragraph in doc.paragraphs:
                paragraph_text = paragraph.text.strip()
                if paragraph_text:
                    content_blocks.append(paragraph_text)

            for table in doc.tables:
                for row in table.rows:
                    row_values = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
                    if row_values:
                        content_blocks.append(" | ".join(row_values))

            extracted_text = "\n\n".join(content_blocks).strip()
            logger.info("Extracted DOCX content blocks=%s from %s", len(content_blocks), docx_path)
            return [(extracted_text, None)]
        except Exception:
            logger.exception("DOCX extraction failed for %s", docx_path)
            raise

    @staticmethod
    def extract_txt(txt_path: str) -> List[Tuple[str, int | None]]:
        encodings = ("utf-8", "utf-8-sig", "latin-1")
        for encoding in encodings:
            try:
                with open(txt_path, "r", encoding=encoding) as file_handle:
                    text = file_handle.read()
                logger.info("Extracted TXT content from %s with encoding %s", txt_path, encoding)
                return [(text.strip(), None)]
            except UnicodeDecodeError:
                continue
            except Exception:
                logger.exception("TXT extraction failed for %s", txt_path)
                raise
        raise ValueError("Unable to decode text file")

    @staticmethod
    def extract(file_path: str, file_type: str) -> List[Tuple[str, int | None]]:
        normalized_type = (file_type or "").lower()
        if normalized_type in {"pdf", "application/pdf"}:
            return TextExtractor.extract_pdf(file_path)
        if normalized_type in {"docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"}:
            return TextExtractor.extract_docx(file_path)
        if normalized_type in {"txt", "text/plain"}:
            return TextExtractor.extract_txt(file_path)
        raise ValueError(f"Unsupported file type: {file_type}")
