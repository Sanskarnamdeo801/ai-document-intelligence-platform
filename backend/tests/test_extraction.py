from pathlib import Path

from docx import Document as DocxDocument

from app.services.text_extractor import TextExtractor


def test_txt_extraction_reads_content(tmp_path):
    txt_path = tmp_path / "sample.txt"
    txt_path.write_text("First line\nSecond line", encoding="utf-8")

    pages = TextExtractor.extract_txt(str(txt_path))

    assert pages
    assert "First line" in pages[0][0]
    assert "Second line" in pages[0][0]


def test_docx_extraction_reads_paragraphs_and_tables(tmp_path):
    docx_path = Path(tmp_path) / "sample.docx"
    document = DocxDocument()
    document.add_paragraph("Contract title")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Party A"
    table.rows[0].cells[1].text = "Party B"
    document.save(docx_path)

    pages = TextExtractor.extract_docx(str(docx_path))

    assert pages
    extracted_text = pages[0][0]
    assert "Contract title" in extracted_text
    assert "Party A" in extracted_text
    assert "Party B" in extracted_text
